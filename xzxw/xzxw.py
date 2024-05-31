# -*- coding = utf-8 -*-
# @Time: 2024/5/29 22:15
# @Author: ZhouYanHui
import os
import requests
import random
import time
import json
from docx import Document


user_agent = [
    'Mozilla/5.0 (compatible; MSIE 8.0; Windows NT 6.0; Trident/4.0; Acoo Browser 1.98.744; .NET CLR 3.5.30729)',
    'Mozilla/5.0 (compatible; MSIE 8.0; Windows NT 6.0; Trident/4.0; Acoo Browser 1.98.744; .NET CLR 3.5.30729)',
    'Mozilla/5.0 (compatible; MSIE 9.0; AOL 9.7; AOLBuild 4343.19; Windows NT 6.1; WOW64; Trident/5.0; FunWebProducts)',
    'Mozilla/5.0 (compatible; MSIE 9.0; AOL 9.1; AOLBuild 4334.5012; Windows NT 6.0; WOW64; Trident/5.0)',
    'Mozilla/5.0 (compatible; MSIE 9.0; AOL 9.0; Windows NT 6.0; Trident/5.0)',
    'Mozilla/5.0 (X11; U; UNICOS lcLinux; en-US) Gecko/20140730 (KHTML, like Gecko, Safari/419.3) Arora/0.8.0',
    'Mozilla/5.0 (X11; U; Linux; ru-RU) AppleWebKit/527+ (KHTML, like Gecko, Safari/419.3) Arora/0.6 (Change: 802 025a17d)',
    'Mozilla/5.0 (X11; U; Linux; pt-PT) AppleWebKit/523.15 (KHTML, like Gecko, Safari/419.3) Arora/0.4'
    'Mozilla/5.0 (X11; U; Linux; en-GB) AppleWebKit/527+ (KHTML, like Gecko, Safari/419.3) Arora/0.3 (Change: 239 52c6958)',
]


def post_data(url, page):
    print(f'正在请求第{page}页:', url)
    # 获取数据
    headers = {
        'User-Agent': random.choice(user_agent),
    }

    data = {
        'q': '非遗',
        'pageNo': page,
        'pageSize': 20,
        'channel': 1,
        'sort': 'date desc',
        'siteID': 5,
        'nodeID': ''
    }

    for _ in range(5):
        try:
            res = requests.post(
                url,
                headers=headers,
                timeout=10,
                data=data
                # proxies=proxies
            )
        except:
            continue
        res.encoding = 'utf-8'
        time.sleep(random.randint(2, 4))
        if res.status_code == 200 and res:
            return res


def save_docx(date, title, content, sourcename):
    doc = Document()
    # 写入标题
    doc.add_heading(title, level=0)
    # 写入来源
    doc.add_paragraph(f'来源：{sourcename}')
    # 写入正文
    doc.add_paragraph(content)
    print(date+title)

    if os.path.exists(date[:4]):
        doc.save(f'{date[:4]}/{date + title}.docx')
    else:
        os.makedirs(date[:4])
        doc.save(f'{date[:4]}/{date + title}.docx')


def xzxw():
    url = 'https://search.xzxw.com/xy/Search.do'
    page = 1
    # page = 230
    while True:
        if page > 235:
            break
        # 发送请求,获取数据
        res = post_data(url, page)
        print(res.text)
        # 解析json文件
        res_list = json.loads(res.text)['article']
        for res_dict in res_list:
            date = res_dict['date']
            # if int(date.replace('-', '')) < 20000000:
            title = res_dict['title']
            enpcontent = res_dict['enpcontent']
            sourcename = res_dict['sourcename']

            save_docx(date, title, enpcontent, sourcename)

        page += 1


def main():
    xzxw()


if __name__ == '__main__':
    main()
